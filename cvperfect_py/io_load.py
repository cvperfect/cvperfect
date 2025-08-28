"""
Moduł do wczytywania CV z różnych formatów (PDF, DOCX, TXT)
Deterministyczny, offline, bez halucynacji
"""

import os
import tempfile
from pathlib import Path
from typing import Dict, Optional, Tuple
import logging

logger = logging.getLogger(__name__)

def load_cv(file_path: str, raw_text: Optional[str] = None) -> Tuple[str, str]:
    """
    Wczytuje CV z pliku lub surowego tekstu.
    
    Args:
        file_path: Ścieżka do pliku CV (.pdf, .docx, .txt) lub None dla raw_text
        raw_text: Opcjonalny surowy tekst (np. z modala)
        
    Returns:
        (content, format): Treść CV i wykryty format
    """
    # Jeśli podano surowy tekst, zapisz jako tymczasowy TXT
    if raw_text:
        logger.info("Processing raw text input")
        return raw_text, "txt"
    
    if not file_path or not os.path.exists(file_path):
        raise FileNotFoundError(f"CV file not found: {file_path}")
    
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    if extension == '.pdf':
        return _load_pdf(file_path), "pdf"
    elif extension == '.docx':
        return _load_docx(file_path), "docx"
    elif extension in ['.txt', '.text']:
        return _load_txt(file_path), "txt"
    else:
        raise ValueError(f"Unsupported file format: {extension}. Use PDF, DOCX or TXT.")


def _load_pdf(file_path: Path) -> str:
    """Wczytuje PDF używając pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
    
    text_parts = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                logger.debug(f"Extracting page {page_num}")
                text = page.extract_text()
                if text:
                    text_parts.append(text)
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        raise ValueError(f"Failed to read PDF: {e}")
    
    return "\n".join(text_parts)


def _load_docx(file_path: Path) -> str:
    """Wczytuje DOCX używając python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    text_parts = []
    try:
        doc = Document(file_path)
        
        # Wyciągnij paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Wyciągnij tabele (flatten to text)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        raise ValueError(f"Failed to read DOCX: {e}")
    
    return "\n".join(text_parts)


def _load_txt(file_path: Path) -> str:
    """Wczytuje zwykły plik tekstowy."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Fallback to latin-1 for older files
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error reading TXT: {e}")
        raise ValueError(f"Failed to read TXT: {e}")


def validate_cv_content(content: str) -> bool:
    """
    Waliduje czy treść wygląda jak CV.
    
    Returns:
        True jeśli treść wygląda OK, False jeśli jest pusta lub zbyt krótka
    """
    if not content:
        return False
    
    # Minimum 100 znaków (bardzo krótkie CV)
    if len(content.strip()) < 100:
        return False
    
    # Sprawdź czy ma jakieś słowa (nie tylko znaki specjalne)
    words = content.split()
    if len(words) < 20:
        return False
    
    return True